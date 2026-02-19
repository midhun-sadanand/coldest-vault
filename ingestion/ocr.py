"""
OCR processing for images and PDFs.
"""

import io
from abc import ABC, abstractmethod
from typing import List, Optional
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes

from config import get_settings

DOCX_MIME_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


class OCREngine(ABC):
    """Abstract base class for OCR engines."""
    
    @abstractmethod
    def extract_text(self, image: Image.Image) -> str:
        """Extract text from an image."""
        pass
    
    def process_images(self, images: List[Image.Image]) -> str:
        """Process multiple images and combine text."""
        texts = []
        for i, image in enumerate(images):
            text = self.extract_text(image)
            if text.strip():
                texts.append(f"--- Page {i + 1} ---\n{text}")
        return "\n\n".join(texts)


class TesseractOCR(OCREngine):
    """Tesseract-based OCR engine (local, free)."""
    
    def __init__(self, lang: str = "eng"):
        self.lang = lang
    
    def extract_text(self, image: Image.Image) -> str:
        """Extract text using Tesseract."""
        try:
            return pytesseract.image_to_string(image, lang=self.lang)
        except Exception as e:
            print(f"  ⚠️ Tesseract error: {e}")
            return ""


class GoogleVisionOCR(OCREngine):
    """Google Cloud Vision OCR engine (cloud, paid but more accurate)."""
    
    def __init__(self):
        try:
            from google.cloud import vision
            self.client = vision.ImageAnnotatorClient()
            self.vision = vision
        except ImportError:
            raise ImportError(
                "google-cloud-vision is required for Google Vision OCR. "
                "Install it with: pip install google-cloud-vision"
            )
    
    def extract_text(self, image: Image.Image) -> str:
        """Extract text using Google Cloud Vision."""
        try:
            # Convert PIL image to bytes
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            content = buffer.getvalue()
            
            vision_image = self.vision.Image(content=content)
            response = self.client.text_detection(image=vision_image)
            
            if response.error.message:
                raise Exception(response.error.message)
            
            texts = response.text_annotations
            if texts:
                return texts[0].description
            return ""
        except Exception as e:
            print(f"  ⚠️ Google Vision error: {e}")
            return ""


class DirectPDFExtractor:
    """Extract text directly from typed/digital PDFs using pdfplumber (no OCR)."""

    def extract_text(self, pdf_bytes: bytes) -> str:
        """Extract embedded text from all pages of a PDF."""
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        texts.append(f"--- Page {i + 1} ---\n{text}")
            return "\n\n".join(texts)
        except Exception as e:
            print(f"  ⚠️ DirectPDFExtractor error: {e}")
            return ""


def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from a Word Document (.docx) file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        print(f"  ⚠️ DOCX extraction error: {e}")
        return ""


def pdf_to_images(pdf_bytes: bytes, max_pages: Optional[int] = None) -> List[Image.Image]:
    """
    Convert PDF bytes to a list of PIL images.
    
    Args:
        pdf_bytes: PDF file content
        max_pages: Maximum number of pages to convert
        
    Returns:
        List of PIL Image objects
    """
    settings = get_settings()
    max_pages = max_pages or settings.max_ocr_pages
    
    images = convert_from_bytes(
        pdf_bytes,
        last_page=max_pages,
        dpi=200  # Good balance between quality and speed
    )
    
    return images


def process_document(
    content: bytes,
    mime_type: str,
    ocr_engine: OCREngine,
    min_direct_chars: int = 200
) -> str:
    """
    Process a document and extract text.

    For typed PDFs, direct text extraction is attempted first; OCR is used as
    a fallback when the embedded text is too short (i.e. the PDF is scanned).
    Word Documents (.docx) are always extracted directly.

    Args:
        content: File content as bytes
        mime_type: MIME type of the file
        ocr_engine: OCR engine to use (fallback for PDFs)
        min_direct_chars: Minimum characters from direct extraction before
                          falling back to OCR for PDFs

    Returns:
        Extracted text
    """
    if mime_type == DOCX_MIME_TYPE:
        return extract_docx_text(content)

    elif mime_type == 'application/pdf':
        direct_text = DirectPDFExtractor().extract_text(content)
        if len(direct_text.strip()) >= min_direct_chars:
            return direct_text
        # Fallback to OCR for scanned / image-only PDFs
        images = pdf_to_images(content)
        return ocr_engine.process_images(images)

    elif mime_type.startswith('image/'):
        image = Image.open(io.BytesIO(content))
        return ocr_engine.extract_text(image)

    else:
        raise ValueError(f"Unsupported mime type: {mime_type}")
